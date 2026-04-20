from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──
for section in doc.sections:
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)

# ── Colour palette ──
DARK   = RGBColor(0x1a, 0x1a, 0x2e)   # dark navy
GOLD   = RGBColor(0xf4, 0xc4, 0x30)   # amber/gold
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
GREY   = RGBColor(0x55, 0x55, 0x55)
LGREY  = RGBColor(0xEE, 0xEE, 0xEE)

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def bold_run(para, text, size=11, color=None, italic=False):
    r = para.add_run(text)
    r.bold  = True
    r.italic = italic
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return r

def normal_run(para, text, size=10, color=None):
    r = para.add_run(text)
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return r

def heading_para(text, size=14, color=DARK, align=WD_ALIGN_PARAGRAPH.LEFT, bold=True, space_before=12, space_after=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'F4C430')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ════════════════════════════════════════════════════
# SECTION 1 – HEADER BANNER
# ════════════════════════════════════════════════════
hdr_tbl = doc.add_table(rows=1, cols=2)
hdr_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_tbl.columns[0].width = Inches(3.5)
hdr_tbl.columns[1].width = Inches(3.5)

left_cell  = hdr_tbl.rows[0].cells[0]
right_cell = hdr_tbl.rows[0].cells[1]
set_cell_bg(left_cell,  '1a1a2e')
set_cell_bg(right_cell, '1a1a2e')

lp = left_cell.paragraphs[0]
lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
bold_run(lp, 'AutoDrive Corp.',  16, WHITE)
lp.add_run('\n')
normal_run(lp, '123 Motor Avenue, Detroit, MI 48201', 9, RGBColor(0xAA,0xAA,0xAA))

rp = right_cell.paragraphs[0]
rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
normal_run(rp, 'Tel: (313) 555-0198\ninfo@autodrivecorp.com\nwww.autodrivecorp.com', 9, RGBColor(0xCC,0xCC,0xCC))

doc.add_paragraph()

# ════════════════════════════════════════════════════
# SECTION 2 – BUSINESS REQUIREMENTS DOCUMENT
# ════════════════════════════════════════════════════
heading_para('Software Development Business Requirements', 15, DARK, space_before=6)
divider()

# Meta grid
meta = [
    ('Project Name:', 'AutoDrive Fleet Management System',  'Date:',        'April 20, 2026'),
    ('Project Manager:', 'Diana Marsh',                    'Department:',  'Information Technology'),
    ('Version:',      '2.1',                               'Status:',      'Active Development'),
]
meta_tbl = doc.add_table(rows=len(meta), cols=4)
meta_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
for row_idx, (k1, v1, k2, v2) in enumerate(meta):
    cells = meta_tbl.rows[row_idx].cells
    cells[0].width = Inches(1.5)
    bold_run(cells[0].paragraphs[0], k1, 10, DARK)
    normal_run(cells[1].paragraphs[0], v1, 10)
    cells[1].width = Inches(2.2)
    bold_run(cells[2].paragraphs[0], k2, 10, DARK)
    normal_run(cells[3].paragraphs[0], v2, 10)
    cells[3].width = Inches(1.5)

doc.add_paragraph()

heading_para('Executive Summary', 11, DARK, space_before=8, space_after=2)
es = doc.add_paragraph()
normal_run(es, (
    'AutoDrive Corp. will develop a comprehensive fleet management and vehicle diagnostics platform '
    'that leverages AI-driven analytics to optimise vehicle performance, predictive maintenance scheduling, '
    'and real-time driver behaviour monitoring. The system integrates IoT telemetry, cloud-based data '
    'storage, and a web dashboard to deliver actionable insights across entire vehicle fleets.'
), 10, GREY)

heading_para('Project Objectives', 11, DARK, space_before=8, space_after=2)
for obj in [
    'Automate real-time telemetry collection from 500+ connected vehicles via IoT sensors.',
    'Deploy AI models to predict component failures with ≥90 % accuracy, reducing unplanned downtime.',
    'Provide fleet managers a live dashboard for vehicle location, fuel consumption, and driver scoring.',
    'Maintain a cloud database of deduplicated, cleaned vehicle records for regulatory compliance.',
]:
    p = doc.add_paragraph(style='List Bullet')
    normal_run(p, obj, 10, GREY)

heading_para('Key Stakeholders', 11, DARK, space_before=8, space_after=4)
sh_data = [
    ('Diana Marsh',   'Project Manager',  'Oversight, timeline, stakeholder communication'),
    ('Carlos Ruiz',   'Lead Developer',   'Core platform: telemetry pipeline, API layer'),
    ('Priya Singh',   'AI/ML Engineer',   'Predictive maintenance models, anomaly detection'),
    ('Tom Walsh',     'QA Engineer',      'Testing, validation, data quality assurance'),
    ('Aiko Tanaka',   'UX Designer',      'Dashboard UI/UX, driver-facing mobile app'),
]
sh_tbl = doc.add_table(rows=1 + len(sh_data), cols=3)
sh_tbl.style = 'Table Grid'
sh_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
for i, hdr in enumerate(['Name', 'Role', 'Responsibility']):
    cell = sh_tbl.rows[0].cells[i]
    set_cell_bg(cell, '1a1a2e')
    bold_run(cell.paragraphs[0], hdr, 10, WHITE)
for r_idx, (name, role, resp) in enumerate(sh_data):
    row = sh_tbl.rows[r_idx + 1]
    normal_run(row.cells[0].paragraphs[0], name, 10)
    normal_run(row.cells[1].paragraphs[0], role, 10)
    normal_run(row.cells[2].paragraphs[0], resp, 10)
    if r_idx % 2 == 1:
        for c in row.cells:
            set_cell_bg(c, 'F9F9F9')

doc.add_page_break()

# ════════════════════════════════════════════════════
# SECTION 3 – PERFORMANCE REVIEW
# ════════════════════════════════════════════════════
# Banner
kpi_tbl = doc.add_table(rows=1, cols=1)
kpi_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
banner_cell = kpi_tbl.rows[0].cells[0]
set_cell_bg(banner_cell, 'F4C430')
bp = banner_cell.paragraphs[0]
bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
bold_run(bp, 'FLEET PERFORMANCE REVIEW', 16, DARK)

doc.add_paragraph()

# Review meta
rv_meta = [
    ('Review Period:', 'Q1 2026',          'Reviewer:',   'Diana Marsh, PM'),
    ('Fleet Size:',    '512 Vehicles',      'Review Date:','April 20, 2026'),
    ('Platform Ver.:', 'v3.2.1',           'Next Review:','July 1, 2026'),
]
rv_tbl = doc.add_table(rows=len(rv_meta), cols=4)
for row_idx, (k1, v1, k2, v2) in enumerate(rv_meta):
    cells = rv_tbl.rows[row_idx].cells
    bold_run(cells[0].paragraphs[0], k1, 10, DARK)
    normal_run(cells[1].paragraphs[0], v1, 10)
    bold_run(cells[2].paragraphs[0], k2, 10, DARK)
    normal_run(cells[3].paragraphs[0], v2, 10)

doc.add_paragraph()

heading_para('KPI Results', 11, DARK, space_before=6, space_after=4)
kpi_data = [
    ('Excellent',         'Predictive Maintenance Accuracy',
     '93.1 % accuracy vs 90 % target. AI model exceeding expectations on engine fault prediction.',
     'On Target'),
    ('Good',              'Fleet Uptime Rate',
     '98.6 % uptime across active vehicles. Two depots achieved 100 % uptime for the quarter.',
     'On Target'),
    ('Satisfactory',      'Average Fuel Efficiency',
     '11.4 L/100 km vs 10.8 target. Route-optimisation module rollout planned for Q2.',
     'Needs Work'),
    ('Needs Improvement', 'Driver Safety Score Compliance',
     'Only 71 % of drivers scoring ≥80. Coaching programme to launch in May 2026.',
     'Below Target'),
]
kpi_res_tbl = doc.add_table(rows=1 + len(kpi_data), cols=4)
kpi_res_tbl.style = 'Table Grid'
for i, hdr in enumerate(['Rating', 'KPI Metric', 'Result', 'Status']):
    cell = kpi_res_tbl.rows[0].cells[i]
    set_cell_bg(cell, 'F4C430')
    bold_run(cell.paragraphs[0], hdr, 10, DARK)
for r_idx, (rating, metric, result, status) in enumerate(kpi_data):
    row = kpi_res_tbl.rows[r_idx + 1]
    bold_run(row.cells[0].paragraphs[0], rating, 10, DARK)
    bold_run(row.cells[1].paragraphs[0], metric, 10, DARK)
    normal_run(row.cells[2].paragraphs[0], result, 10, GREY)
    normal_run(row.cells[3].paragraphs[0], status, 10, GREY)
    if r_idx % 2 == 1:
        for c in row.cells:
            set_cell_bg(c, 'F0F0F0')

doc.add_page_break()

# ════════════════════════════════════════════════════
# SECTION 4 – SIX BENEFITS
# ════════════════════════════════════════════════════
ben_banner = doc.add_table(rows=1, cols=1)
ben_banner.alignment = WD_TABLE_ALIGNMENT.CENTER
ben_cell = ben_banner.rows[0].cells[0]
set_cell_bg(ben_cell, 'F4C430')
bep = ben_cell.paragraphs[0]
bep.alignment = WD_ALIGN_PARAGRAPH.CENTER
bold_run(bep, 'SIX BENEFITS OF AI-POWERED FLEET MANAGEMENT', 15, DARK)
bep.add_run('\n')
normal_run(bep, 'How AutoDrive transforms vehicle operations for modern businesses', 10, DARK)

doc.add_paragraph()

benefits = [
    ('1', 'Reduced Downtime',
     'Predictive maintenance alerts flag component failures weeks in advance, cutting unplanned downtime by up to 40 %.'),
    ('2', 'Lower Fuel Costs',
     'AI route optimisation and eco-driving coaching reduce average fuel consumption by 12–18 % per vehicle per year.'),
    ('3', 'Enhanced Driver Safety',
     'Real-time behaviour scoring (harsh braking, speeding, distraction) reduces accident rates and lowers insurance premiums.'),
    ('4', 'Full Fleet Visibility',
     'Live GPS tracking and status dashboards give managers instant awareness of every vehicle\'s location and condition.'),
    ('5', 'Regulatory Compliance',
     'Automated digital logbooks and tachograph reports keep fleets compliant with transport regulations without manual paperwork.'),
    ('6', 'Scalability',
     'Cloud-native architecture supports fleets from 10 to 10,000 vehicles with no infrastructure changes required.'),
]

ben_tbl = doc.add_table(rows=3, cols=2)
ben_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
ben_tbl.style = 'Table Grid'
for idx, (num, title, desc) in enumerate(benefits):
    row = idx // 2
    col = idx % 2
    cell = ben_tbl.rows[row].cells[col]
    p = cell.paragraphs[0]
    bold_run(p, f'{num}  ', 20, GOLD)
    bold_run(p, title + '\n', 11, DARK)
    normal_run(p, desc, 10, GREY)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)

doc.add_page_break()

# ════════════════════════════════════════════════════
# SECTION 5 – SEVEN TIPS
# ════════════════════════════════════════════════════
tip_banner = doc.add_table(rows=1, cols=1)
tip_banner.alignment = WD_TABLE_ALIGNMENT.CENTER
tip_cell = tip_banner.rows[0].cells[0]
set_cell_bg(tip_cell, '1a1a2e')
tp = tip_cell.paragraphs[0]
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
bold_run(tp, '7 TIPS TO MAXIMISE YOUR FLEET\'S PERFORMANCE', 15, GOLD)
tp.add_run('\n')
normal_run(tp, 'Follow these best practices to get the most out of the AutoDrive platform', 10, WHITE)

doc.add_paragraph()

tips = [
    ('01', 'Configure Alerts Correctly',
     'Set maintenance alert thresholds specific to each vehicle make and model to avoid alert fatigue and missed warnings.'),
    ('02', 'Review Driver Scores Weekly',
     'Weekly coaching sessions based on driver score reports produce the fastest improvements in safety and fuel efficiency.'),
    ('03', 'Use Route Optimisation Daily',
     'Enable the AI route planner every morning. Dynamic rerouting around traffic alone saves an average of 22 minutes per vehicle per day.'),
    ('04', 'Keep Vehicle Profiles Updated',
     'Accurate tyre sizes, engine specs, and service histories improve the accuracy of predictive maintenance models by up to 15 %.'),
    ('05', 'Integrate Fuel Card Data',
     'Sync your fuel card provider to cross-reference actual fill-ups against telematics consumption data and catch fuel theft or inefficiencies.'),
    ('06', 'Schedule Preventive Services Early',
     'Book vehicles in for service as soon as an alert fires — not when they\'re already symptomatic — to avoid costly breakdowns.'),
    ('07', 'Export Monthly Reports',
     'Use the one-click compliance export to generate tachograph, DVLA, and emissions reports automatically at month-end.'),
]

for num, title, desc in tips:
    tip_tbl = doc.add_table(rows=1, cols=2)
    tip_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    num_cell  = tip_tbl.rows[0].cells[0]
    desc_cell = tip_tbl.rows[0].cells[1]
    num_cell.width  = Inches(0.7)
    desc_cell.width = Inches(6.3)
    set_cell_bg(num_cell, 'F4C430')
    np2 = num_cell.paragraphs[0]
    np2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bold_run(np2, num, 14, DARK)
    dp = desc_cell.paragraphs[0]
    bold_run(dp, title + '  ', 11, DARK)
    normal_run(dp, desc, 10, GREY)
    dp.paragraph_format.space_before = Pt(4)
    dp.paragraph_format.space_after  = Pt(4)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ════════════════════════════════════════════════════
# FOOTER
# ════════════════════════════════════════════════════
divider()
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
normal_run(footer_p, '© 2026 AutoDrive Corp.  ·  All rights reserved  ·  info@autodrivecorp.com', 9, GREY)

out = '/home/user/HunterAI/AutoDrive_Business_Requirements.docx'
doc.save(out)
print(f'Saved: {out}')
